from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any


@dataclass(frozen=True)
class ParsedPage:
    page_number: int
    text: str
    width: float | None = None
    height: float | None = None
    ocr_status: str = "not_required"
    blocks: list[dict[str, Any]] | None = None


@dataclass(frozen=True)
class ParsedDocument:
    filename: str
    content_type: str
    text: str
    pages: list[ParsedPage]


def parse_document(filename: str, data: bytes, content_type: str | None = None) -> ParsedDocument:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        return _parse_pdf(filename, data)
    if extension == ".docx":
        return _parse_docx(filename, data)
    if extension == ".md":
        text = _decode_text(data)
        return _single_page(filename, "text/markdown", _markdown_to_text(text))
    if extension == ".html":
        text = _html_to_text(data.decode("utf-8", errors="ignore"))
        return _single_page(filename, "text/html", text)
    text = _decode_text(data)
    return _single_page(filename, content_type or "text/plain", text)


def parse_html(filename: str, html: str) -> ParsedDocument:
    return _single_page(filename, "text/html", _html_to_text(html))


def locate_chunk(chunk_text: str, pages: list[ParsedPage]) -> dict[str, Any]:
    normalized_chunk = _compact(chunk_text[:240])
    for page in pages:
        if normalized_chunk and normalized_chunk in _compact(page.text):
            return {
                "page_number": page.page_number,
                "bbox": _first_useful_bbox(page.blocks or [], page.width, page.height),
            }
    if pages:
        page = pages[0]
        return {"page_number": page.page_number, "bbox": _first_useful_bbox(page.blocks or [], page.width, page.height)}
    return {"page_number": 1, "bbox": {}}


def _parse_pdf(filename: str, data: bytes) -> ParsedDocument:
    import fitz

    pages: list[ParsedPage] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for index, page in enumerate(doc, start=1):
            blocks = _pdf_blocks(page)
            text = "\n".join(block["text"] for block in blocks if block["text"]).strip()
            ocr_status = "not_required"
            if not text:
                text = _ocr_pdf_page(page)
                ocr_status = "completed"
                blocks = [{"text": text, "bbox": {"x0": 0.0, "y0": 0.0, "x1": float(page.rect.width), "y1": float(page.rect.height)}}]
            pages.append(
                ParsedPage(
                    page_number=index,
                    text=text,
                    width=float(page.rect.width),
                    height=float(page.rect.height),
                    ocr_status=ocr_status,
                    blocks=blocks,
                )
            )
    full_text = "\n\n".join(page.text for page in pages if page.text).strip()
    if not full_text:
        raise ValueError("PDF contains no extractable text after OCR")
    return ParsedDocument(filename=filename, content_type="application/pdf", text=full_text, pages=pages)


def _pdf_blocks(page: Any) -> list[dict[str, Any]]:
    result: list[dict[str, Any]] = []
    for block in page.get_text("blocks"):
        if len(block) < 5:
            continue
        x0, y0, x1, y1, text = block[:5]
        value = str(text).strip()
        if not value:
            continue
        result.append(
            {
                "text": value,
                "bbox": {"x0": float(x0), "y0": float(y0), "x1": float(x1), "y1": float(y1)},
            }
        )
    return result


def _ocr_pdf_page(page: Any) -> str:
    import fitz
    import pytesseract
    from PIL import Image

    pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2))
    image = Image.open(BytesIO(pixmap.tobytes("png")))
    try:
        text = pytesseract.image_to_string(image, lang="chi_sim+eng")
    except Exception:  # noqa: BLE001
        text = pytesseract.image_to_string(image, lang="eng")
    if not text.strip():
        raise ValueError(f"OCR produced no text for PDF page {page.number + 1}")
    return text.strip()


def _parse_docx(filename: str, data: bytes) -> ParsedDocument:
    from docx import Document as DocxDocument

    doc = DocxDocument(BytesIO(data))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    return _single_page(filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "\n\n".join(paragraphs))


def _single_page(filename: str, content_type: str, text: str) -> ParsedDocument:
    text = text.strip()
    if not text:
        raise ValueError("Document contains no extractable text")
    page = ParsedPage(
        page_number=1,
        text=text,
        ocr_status="not_required",
        blocks=[{"text": text[:1000], "bbox": {"x0": 0.0, "y0": 0.0, "x1": 100.0, "y1": 100.0}}],
    )
    return ParsedDocument(filename=filename, content_type=content_type, text=text, pages=[page])


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _markdown_to_text(markdown: str) -> str:
    try:
        from markdown_it import MarkdownIt
        from bs4 import BeautifulSoup

        html = MarkdownIt().render(markdown)
        return BeautifulSoup(html, "html.parser").get_text("\n")
    except Exception:  # noqa: BLE001
        return markdown


def _html_to_text(html: str) -> str:
    try:
        from readability import Document

        html = Document(html).summary()
    except Exception:  # noqa: BLE001
        pass
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, "html.parser")
    for node in soup(["script", "style", "noscript", "svg"]):
        node.decompose()
    return re.sub(r"\n{3,}", "\n\n", soup.get_text("\n")).strip()


def _compact(value: str) -> str:
    return re.sub(r"\s+", "", value)


def _first_useful_bbox(blocks: list[dict[str, Any]], width: float | None, height: float | None) -> dict[str, float]:
    for block in blocks:
        bbox = block.get("bbox") if isinstance(block, dict) else None
        if bbox:
            return {key: float(value) for key, value in bbox.items()}
    return {"x0": 0.0, "y0": 0.0, "x1": float(width or 100.0), "y1": float(height or 100.0)}
