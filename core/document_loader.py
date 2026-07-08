from __future__ import annotations

import hashlib
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config.settings import AppSettings

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _read_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(page.strip() for page in pages if page.strip())


def _read_docx(data: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(BytesIO(data))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    return "\n\n".join(paragraphs)


def extract_text(filename: str, data: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported document type: {extension or '<none>'}. Supported: {supported}")
    if extension in {".txt", ".md"}:
        return _decode_text(data)
    if extension == ".pdf":
        return _read_pdf(data)
    if extension == ".docx":
        return _read_docx(data)
    raise ValueError(f"Unsupported document type: {extension}")


def hash_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def normalize_text_for_hash(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    lines = [line.rstrip() for line in normalized.split("\n")]
    return "\n".join(lines).strip()


def hash_text(text: str) -> str:
    return hashlib.sha256(normalize_text_for_hash(text).encode("utf-8")).hexdigest()


def build_chunks(
    filename: str,
    content_type: str | None,
    data: bytes,
    settings: AppSettings,
    document_id: str | None = None,
    content_sha256: str | None = None,
    user_id: str = "default",
) -> tuple[str, list[Document]]:
    text = extract_text(filename, data).strip()
    resolved_document_id = document_id or str(uuid4())
    resolved_content_sha256 = content_sha256 or hash_text(text)
    return build_chunks_from_text(
        filename=filename,
        content_type=content_type,
        text=text,
        settings=settings,
        document_id=resolved_document_id,
        content_sha256=resolved_content_sha256,
        user_id=user_id,
    )


def build_chunks_from_text(
    filename: str,
    content_type: str | None,
    text: str,
    settings: AppSettings,
    document_id: str,
    content_sha256: str,
    user_id: str,
) -> tuple[str, list[Document]]:
    text = text.strip()
    if not text:
        raise ValueError("Document contains no extractable text")

    upload_time = datetime.now(timezone.utc).isoformat()
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", "。", ".", " ", ""],
    )
    chunks = splitter.split_text(text)
    documents = [
        Document(
            page_content=chunk,
            metadata={
                "document_id": document_id,
                "file_id": document_id,
                "user_id": user_id,
                "filename": filename,
                "content_type": content_type or "application/octet-stream",
                "content_sha256": content_sha256,
                "chunk_index": index,
                "upload_time": upload_time,
            },
        )
        for index, chunk in enumerate(chunks)
    ]
    return document_id, documents
